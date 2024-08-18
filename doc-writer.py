from docx import Document
import json
from common import set_margin, stylize
import datetime


with open("./profiledata.json", "r") as fp:
    profile_data = json.loads(fp.read())
# Create a new Document
doc = Document()


# Access the section element of the document to modify the page setup
section = doc.sections[0]

# Apply the 'pgMar' element to the document
pgMar = set_margin()
section._sectPr.append(pgMar)

# Add header
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = profile_data["basic"]["name"]

# Set the font, size, and bold properties
title = header_paragraph.runs[0]
stylize(title, bold=True)


# Add a second line in the header with different formatting
email = header_paragraph.add_run(
    f"\n{profile_data['basic']['email']} | {profile_data['basic']['contact']} | {profile_data['basic']['social']}"
)
stylize(email, size=10)

# Add the Skill Set Section
section_heading = doc.add_paragraph()
section = section_heading.add_run("Skill Set")
stylize(section, size=14, bold=True)

skillset = profile_data["skillset"].keys()

for skill in skillset:

    # Add the "Skill" line with specific formatting
    skill_run_bold_italic = section_heading.add_run(f"\n{skill.capitalize()}:\t  ")
    stylize(skill_run_bold_italic, size=10, bold=True, italic=True)

    skill_list = " ".join(profile_data["skillset"][skill])

    # Add the rest of the line with normal formatting
    skill_run_normal = section_heading.add_run(skill_list)
    stylize(skill_run_normal, size=10)

# Add the Experience Section
section = section_heading.add_run("\nExperience")
stylize(section, size=14, bold=True)


for idx, experience in enumerate(profile_data["experiences"]):

    # Add the "Company Name" line with specific formatting
    company_run_bold_italic = section_heading.add_run(
        f"\n{idx + 1}. {experience['companyName']}"
    )
    stylize(company_run_bold_italic, size=12, bold=True, italic=True)

    # Add the "Job Titles" line with specific formatting
    for position in experience["positions"]:
        company_run_bold_italic = section_heading.add_run(
            f"\n{position['title']}" + "\t"
        )
        stylize(company_run_bold_italic, size=10, bold=True)
        run_right = section_heading.add_run(
            f"{position['tenure']['from']} - {position['tenure']['to']}"
        )
        stylize(run_right, size=10, bold=True, italic=True)

    # Add Descriptions
    for desc in experience["description"]:
        desc_run_normal = section_heading.add_run(f"\n\u2022 {desc}")
        stylize(desc_run_normal, size=10)

# Add the Education Section
section = section_heading.add_run("\nEducation")
stylize(section, size=14, bold=True)

for idx, education in enumerate(profile_data["educations"]):
    # Add the "Degree Name" line with specific formatting
    degree_run_bold_italic = section_heading.add_run(
        f"\n{idx + 1}. {education['title']}" + "\t"
    )
    stylize(degree_run_bold_italic, size=10, italic=True)

    run_right = section_heading.add_run(f"{education['grade']}")
    stylize(run_right, size=10, bold=True, italic=True)

# Add the Certifications Section
section = section_heading.add_run("\nCertifications")
stylize(section, size=14, bold=True)

for cert in profile_data["certifications"]:
    cert_run_normal = section_heading.add_run(f"\n\u2022 {cert['title']}")
    stylize(cert_run_normal, size=10)

# Add the SoftSkills Section
section = section_heading.add_run("\nSoft Skills")
stylize(section, size=14, bold=True)

for skill in profile_data["softskills"]:
    company_run_bold = section_heading.add_run(f"\n{skill['title']}" + "\t")
    stylize(company_run_bold, size=10, bold=True)

    run_right = section_heading.add_run(
        f"{skill['tenure']['from']} - {skill['tenure']['to']}"
    )
    stylize(run_right, size=10, bold=True, italic=True)

    # Add Descriptions
    for desc in skill["description"]:
        desc_run_normal = section_heading.add_run(f"\n\u2022 {desc}")
        stylize(desc_run_normal, size=10)

# Save the document
mydate = datetime.datetime.now().strftime("%B-%Y")
doc.save(f'{profile_data["basic"]["name"].lower().replace(" ","")}-resume.docx')
